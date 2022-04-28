import nltk.tokenize
from nltk import pos_tag, word_tokenize, RegexpParser
from string import punctuation
import docx
import pickle

import Vocabulary
import Sentence

chunker = RegexpParser("""
                       NP: {<DT>?<JJ>*<NN>}    #To extract Noun Phrases
                       P: {<IN>}               #To extract Prepositions
                       V: {<V.*>}              #To extract Verbs
                       PP: {<p> <NP>}          #To extract Prepositional Phrases
                       VP: {<V> <NP|PP>*}      #To extract Verb Phrases
                       """)


def read_text_from_file(filename):
    doc = docx.Document(filename)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)


def write_text_to_file(filename, voc):
    doc_file = docx.Document()
    doc_file.add_heading("Предложения", 1)
    for sentence in dict(voc.get_all_sentences()).values():
        doc_file.add_paragraph("---------------------------------------------------------------------------------")
        sentence_text = "Предложение: " + "\'" + str(sentence.get_string()) + "\"" + "\n дерево -> " + str(
            sentence.get_tree())
        doc_file.add_paragraph(sentence_text)
    doc_file.save(filename + ".docx")


def write_vocabulary_to_file(voc, filename):
    with open(filename + ".pkl", "wb") as file:
        result_list = []
        for sent in dict(voc.get_all_sentences()).values():
            sent_list = [sent.get_string(), sent.get_elements(), sent.get_tugged()]
            result_list.append(sent_list)
        pickle.dump(result_list, file)


def read_vocabulary_from_file(filename):
    with open(filename, "rb") as file:
        voc = Vocabulary.Vocabulary()
        result_list = list(pickle.load(file))
        for item in result_list:
            s = Sentence.Sentence()
            s.set_string(item[0])
            s.set_elements(list(item[1]))
            tagged = item[2]
            s.set_tugged(tagged)
            s.set_tree(chunker.parse(tagged))
            voc.add_sentence(s)
        return voc


def process_text(text):
    voc = Vocabulary.Vocabulary()

    sentence_list = list(nltk.tokenize.sent_tokenize(text))
    for sent in sentence_list:
        s = process_sentence(sent)
        voc.add_sentence(s)

    return voc


def process_sentence(sent):
    s = Sentence.Sentence()
    tagged = pos_tag(list(filter(lambda val: val not in punctuation, word_tokenize(sent))))
    elements = [item[1] for item in tagged]
    tree = chunker.parse(tagged)
    s.set_string(sent)
    s.set_elements(elements)
    s.set_tree(tree)
    s.set_tugged(tagged)
    return s
